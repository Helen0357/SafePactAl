"""
DOCX upload support — extraction + /api/contracts/analyze accepts PDF, DOCX, TXT,
and rejects everything else. Orchestrator (Gemini) is mocked — no quota used.
"""
import io
import sys
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

_AGENT_ROOT = Path(__file__).resolve().parent.parent.parent / "agent"
if str(_AGENT_ROOT) not in sys.path:
    sys.path.insert(0, str(_AGENT_ROOT))

from fastapi.testclient import TestClient

from app.main import app

client = TestClient(app)

REPORT = {
    "contract_type": "Rental Agreement",
    "overall_risk": "High",
    "final_recommendation": "Do Not Sign Yet",
    "summary": "Risky.",
    "confidence": 0.9,
    "risks": [
        {"id": "risk_001", "title": "Late Fee", "severity": "High",
         "category": "Fees", "clause_text": "10% daily late fee.",
         "simple_explanation": "High late fee.", "why_it_matters": "Costly.",
         "question_to_ask": "Cap it?", "suggested_action": "Negotiate"},
    ],
    "missing_information": [], "recommended_questions": [],
}

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _make_docx_bytes(paragraphs):
    import docx
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_pdf_bytes(text):
    import fitz  # PyMuPDF
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    return doc.tobytes()


def _patch_orchestrator():
    mock_orch = MagicMock()
    mock_orch.analyze_contract = AsyncMock(return_value=REPORT)
    return patch("app.services.contract_service._get_orchestrator", return_value=mock_orch), mock_orch


# ── DOCX extraction unit ───────────────────────────────────────────────────────

class TestDocxExtraction:
    def test_extracts_paragraphs(self):
        from app.utils.file_utils import extract_text_from_docx
        text = extract_text_from_docx(_make_docx_bytes(
            ["This is a rental agreement.", "Tenant pays a 10% daily late fee."]
        ))
        assert "rental agreement" in text
        assert "late fee" in text

    def test_extracts_table_cells(self):
        import docx
        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Deposit"
        table.rows[0].cells[1].text = "Non-refundable"
        buf = io.BytesIO()
        document.save(buf)
        from app.utils.file_utils import extract_text_from_docx
        text = extract_text_from_docx(buf.getvalue())
        assert "Deposit" in text and "Non-refundable" in text

    def test_empty_docx_raises(self):
        from app.utils.file_utils import extract_text_from_docx
        with pytest.raises(ValueError):
            extract_text_from_docx(_make_docx_bytes([]))

    def test_invalid_bytes_raise(self):
        from app.utils.file_utils import extract_text_from_docx
        with pytest.raises(ValueError):
            extract_text_from_docx(b"this is not a docx file at all")


# ── /api/contracts/analyze upload routing ──────────────────────────────────────

class TestUploadRouting:
    def test_docx_upload_accepted(self):
        docx_bytes = _make_docx_bytes(
            ["This is a rental agreement.", "Tenant pays a 10% daily late fee."]
        )
        p, mock_orch = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("contract.docx", docx_bytes, _DOCX_MIME)},
            )
        assert resp.status_code == 200, resp.text
        assert "session_id" in resp.json()
        # The extracted DOCX text reached the analysis pipeline.
        sent_text = mock_orch.analyze_contract.call_args.args[0]
        assert "late fee" in sent_text.lower()

    def test_docx_language_header_passed(self):
        docx_bytes = _make_docx_bytes(["عقد إيجار مع شروط مجحفة."])
        p, mock_orch = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("contract.docx", docx_bytes, _DOCX_MIME)},
                headers={"X-Language": "ar"},
            )
        assert resp.status_code == 200
        assert mock_orch.analyze_contract.call_args.kwargs.get("language") == "ar"

    def test_pdf_upload_still_works(self):
        pdf_bytes = _make_pdf_bytes("This is a rental agreement with a late fee.")
        p, mock_orch = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("contract.pdf", pdf_bytes, "application/pdf")},
            )
        assert resp.status_code == 200, resp.text
        assert "rental agreement" in mock_orch.analyze_contract.call_args.args[0].lower()

    def test_txt_upload_still_works(self):
        p, mock_orch = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("contract.txt", b"This is a rental agreement.", "text/plain")},
            )
        assert resp.status_code == 200, resp.text
        assert "rental agreement" in mock_orch.analyze_contract.call_args.args[0].lower()

    def test_unsupported_file_rejected(self):
        p, _ = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("photo.jpg", b"\xff\xd8\xff\xe0 not a contract", "image/jpeg")},
            )
        assert resp.status_code == 400
        assert "Unsupported file type" in resp.json()["error"]

    def test_empty_file_rejected(self):
        p, _ = _patch_orchestrator()
        with p:
            resp = client.post(
                "/api/contracts/analyze",
                files={"file": ("contract.docx", b"", _DOCX_MIME)},
            )
        assert resp.status_code == 400
